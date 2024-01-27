import os
import io
import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
from getpass import getpass
from bs4 import BeautifulSoup
import time
from docx import Document
from docx.shared import Cm

def udemy_login():
    driver = create_driver()
    driver.get('https://www.udemy.com/join/login-popup/')
    time.sleep(3)

    # Get user input for email and password
    email_input = input("Enter your Udemy email: ")
    password_input = getpass("Enter your Udemy password: ")

    # Find the email and password fields and fill them with user input
    email = driver.find_element(By.XPATH, '//*[@id="form-group--1"]')
    password = driver.find_element(By.XPATH, '//*[@id="form-group--3"]')

    email.send_keys(email_input)
    password.send_keys(password_input)

    # Now, click the login button
    driver.find_element(By.XPATH, '//*[@id="udemy"]/div[1]/div[2]/div/div/form/button').click()

def create_driver():
    # Specify the full path to the geckodriver.exe
    gecko_path = 'C:\\Users\\Acer\\Desktop\\Udemy\\geckodriver.exe'

    # Create a FirefoxOptions instance
    firefox_options = webdriver.FirefoxOptions()

    # Specify the full path to the Firefox binary
    firefox_binary_path = 'C:\\Program Files\\Mozilla Firefox\\firefox.exe'
    firefox_options.binary_location = firefox_binary_path

    # Create a Firefox driver instance
    return webdriver.Firefox(options=firefox_options)


def save_course_info_to_txt(course_name, sub_topic, price, rate):
    # Create a directory with the course_name
    directory_name = f'({course_name})-Template'
    os.makedirs(directory_name, exist_ok=True)

    # Construct the file path within the new directory
    file_path = os.path.join(directory_name, '00-readme.txt')

    # Write the course information to the file
    with open(file_path, 'w') as f:
        f.write(f'Course Name : {course_name}\n')
        f.write(f'Sub Topic : {sub_topic}\n')
        f.write(f'Price : {price}\n')
        f.write(f'Rate : {rate}\n')

    print('File saved')


def save_screenshot_and_doc(driver, course_name, xpath, paragraph_name, output_directory):
    # Create a Word document
    doc = Document()

    # Add the course information
    doc.add_paragraph(f'{paragraph_name}: {driver.find_element(By.XPATH, xpath).text}')

    # Save the screenshot
    screenshot = driver.find_element(By.XPATH, xpath).screenshot_as_png
    image_stream = io.BytesIO(screenshot)
    doc.add_picture(image_stream, width=Cm(10))  # You may need to adjust the width

    # Save the Word document
    doc_path = os.path.join(output_directory,
                            f'{course_name.replace(" ", "_")}_{paragraph_name.replace(" ", "_").lower()}_info.docx')
    doc.save(doc_path)

    print(f'Doc file with screenshot saved for {course_name} - {paragraph_name}')


def main():
    # Specify the existing directory to save the DOCX files
    output_directory = 'PythonForBeginnersIntro(free)'
    os.makedirs(output_directory, exist_ok=True)

    driver = create_driver()

    # Call the function to execute the login process
    # udemy_login()

    # Navigate to the desired URL
    driver.get('https://www.udemy.com/course/pythonforbeginnersintro/')
    time.sleep(2)

    # Use find_element with By.XPATH to get a specific button by XPath
    driver.find_element(By.XPATH,
                        '//button[@class="ud-btn ud-btn-large ud-btn-ghost ud-heading-md ud-nav-button tabs-module--nav-button--3RilJ"]').click()
    time.sleep(5)

    # Get the page source
    html_doc = driver.page_source

    soup = BeautifulSoup(html_doc, 'html.parser')
    sections = soup.find_all('div', class_='heading')

    for section in sections:
        course_name = section.find('h1', class_='ud-heading-xl clp-lead__title clp-lead__title--small')
        sub_topic = section.find('div', class_='ud-text-md clp-lead__headline')
        price = section.find('div', class_='ud-badge ud-heading-xs course-badges-module--free--1PepP')
        rate = section.find('div', class_='styles--rating-wrapper--ajCRv')

        if course_name:
            course_name = course_name.text.strip()
        else:
            course_name = "N/A"

        if sub_topic:
            sub_topic = sub_topic.text.strip()
        else:
            sub_topic = "N/A"

        if price:
            price = price.text.strip()
        else:
            price = "N/A"

        if rate:
            rate = rate.text.strip()
        else:
            rate = "N/A"

        save_course_info_to_txt(course_name, sub_topic, price, rate)

        save_screenshot_and_doc(driver, course_name,
                                '/html/body/div[1]/div[2]/div/div/div[2]/div[3]/div[2]/div[1]/div/h1', 'Course Name',
                                output_directory)
        save_screenshot_and_doc(driver, course_name,
                                '/html/body/div[1]/div[2]/div/div/div[2]/div[3]/div[2]/div[1]/div/div[1]', 'Sub Topic',
                                output_directory)
        save_screenshot_and_doc(driver, course_name,
                                '/html/body/div[1]/div[2]/div/div/div[2]/div[3]/div[2]/div[1]/div/div[2]/div[1]',
                                'Price', output_directory)
        save_screenshot_and_doc(driver, course_name,
                                '/html/body/div[1]/div[2]/div/div/div[2]/div[3]/div[2]/div[1]/div/div[2]/div[2]',
                                'Rate', output_directory)

    driver.quit()


if __name__ == "__main__":
    main()



